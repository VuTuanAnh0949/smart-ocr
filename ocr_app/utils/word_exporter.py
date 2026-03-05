"""
Word Export Module - Xuất kết quả OCR ra Word với giữ nguyên format
Author: Vũ Tuấn Anh
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io

class WordExporter:
    """Xuất OCR results ra Word file với format preservation"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document_style()
    
    def _setup_document_style(self):
        """Cấu hình style mặc định cho document"""
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def add_title(self, title_text):
        """Thêm tiêu đề chính"""
        title = self.doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return title
    
    def add_metadata(self, metadata):
        """Thêm metadata (ngày, tác giả, etc)"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for key, value in metadata.items():
            run = p.add_run(f"{key}: {value}\n")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_original_image(self, image, max_width=6.0):
        """Thêm ảnh gốc vào document"""
        self.doc.add_heading('Ảnh gốc:', level=2)
        
        # Convert PIL Image to bytes
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        # Add image to document
        self.doc.add_picture(img_byte_arr, width=Inches(max_width))
        self.doc.add_paragraph()  # Add spacing
    
    def add_extracted_text(self, text, preserve_layout=True):
        """Thêm text đã extract với/không preserve layout"""
        self.doc.add_heading('Văn bản trích xuất:', level=2)
        
        if preserve_layout:
            # Preserve layout: dùng font monospace và giữ nguyên format
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Normal text
            self.doc.add_paragraph(text)
    
    def add_table_from_text(self, table_data):
        """Thêm bảng từ data đã phân tích"""
        if not table_data or len(table_data) == 0:
            return
        
        # Create table
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = str(cell_text)
    
    def add_page_break(self):
        """Thêm page break"""
        self.doc.add_page_break()
    
    def add_section(self, title, content):
        """Thêm một section với title và content"""
        self.doc.add_heading(title, level=2)
        self.doc.add_paragraph(content)
    
    def export_pdf_pages(self, pdf_pages_data):
        """Xuất nhiều trang PDF với format giữ nguyên"""
        for i, page_data in enumerate(pdf_pages_data):
            self.doc.add_heading(f'Trang {i + 1}', level=2)
            
            # Add image if available
            if 'image' in page_data:
                img_byte_arr = io.BytesIO()
                page_data['image'].save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)
                self.doc.add_picture(img_byte_arr, width=Inches(6.0))
            
            # Add extracted text
            if 'text' in page_data:
                p = self.doc.add_paragraph()
                run = p.add_run(page_data['text'])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            # Add page break except for last page
            if i < len(pdf_pages_data) - 1:
                self.add_page_break()
    
    def save(self, output_path):
        """Lưu document"""
        self.doc.save(output_path)
        return output_path
    
    def save_to_bytes(self):
        """Lưu document ra bytes buffer (cho download)"""
        byte_io = io.BytesIO()
        self.doc.save(byte_io)
        byte_io.seek(0)
        return byte_io.getvalue()


def create_word_from_ocr_result(image, text, preserve_layout=True, metadata=None):
    """
    Tạo Word document từ kết quả OCR
    
    Args:
        image: PIL Image object
        text: Extracted text
        preserve_layout: Giữ nguyên layout hay không
        metadata: Dict chứa metadata (optional)
    
    Returns:
        bytes: Word document content
    """
    exporter = WordExporter()
    
    # Add title
    exporter.add_title('Kết quả OCR')
    
    # Add metadata
    if metadata:
        exporter.add_metadata(metadata)
    
    # Add original image
    exporter.add_original_image(image)
    
    # Add extracted text
    exporter.add_extracted_text(text, preserve_layout)
    
    # Return as bytes
    return exporter.save_to_bytes()


def create_word_from_pdf_ocr(pdf_pages_data, metadata=None):
    """
    Tạo Word document từ PDF OCR results
    
    Args:
        pdf_pages_data: List of dict với keys 'image' và 'text' cho mỗi page
        metadata: Dict chứa metadata (optional)
    
    Returns:
        bytes: Word document content
    """
    exporter = WordExporter()
    
    # Add title
    exporter.add_title('Kết quả OCR - PDF')
    
    # Add metadata
    if metadata:
        exporter.add_metadata(metadata)
    
    # Export all pages
    exporter.export_pdf_pages(pdf_pages_data)
    
    # Return as bytes
    return exporter.save_to_bytes()
